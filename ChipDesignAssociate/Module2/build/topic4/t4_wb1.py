# -*- coding: utf-8 -*-
"""Topic 4 workbook — front matter + Part 1: the Verilog language."""
import _boot
from wbkit import *
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL


def B(t, d=None, **kw):
    kw.update(d or {}); kw["b"] = True; return (t, kw)


def N(t, d=None, **kw):
    kw.update(d or {}); return (t, kw)


def I(t, d=None, **kw):
    kw.update(d or {}); kw["i"] = True; return (t, kw)


def M(t, d=None, **kw):
    kw.update(d or {}); kw["f"] = MONOF; return (t, kw)


def build(w):
    # ------------------------------------------------------------ cover
    w.para([N("CHIP DESIGN ASSOCIATE  ·  O-LEVEL ‘CHIP DESIGN’",
              {"b": True, "s": 11, "c": TEAL})], space_after=2)
    p = w.d.add_paragraph()
    r = p.add_run("Module 2 — Topic 4")
    r.font.name = HEADF; r.font.size = Pt(15); r.font.bold = True; r.font.color.rgb = SLATE
    p.paragraph_format.space_after = Pt(2)
    p = w.d.add_paragraph()
    r = p.add_run("RTL Design Using HDL")
    r.font.name = HEADF; r.font.size = Pt(25); r.font.bold = True; r.font.color.rgb = NAVY
    p.paragraph_format.space_after = Pt(4)
    p = w.d.add_paragraph()
    r = p.add_run("Tutorial & Practice Workbook")
    r.font.name = HEADF; r.font.size = Pt(16); r.font.color.rgb = AMBER
    p.paragraph_format.space_after = Pt(10)
    w.para([N("A self-study companion to the Topic 4 slide deck. It elaborates every construct "
              "the deck introduces, explains what each one becomes in hardware and why, walks "
              "through six guided tutorials at the keyboard, and ends with 60 graded exercises "
              "and full worked solutions. Every design referred to here exists as verified, "
              "runnable code in Topic4_Lab/. Nothing in this workbook requires you to look "
              "anything up elsewhere.", {"s": 10.5})])
    w.para([N("NOS: NIE/ELE/N0102  ·  Module 2 “Verilog RTL coding for Synthesis”, Topic 4  ·  "
              "Syllabus: introduction to Verilog syntax and constructs; designing combinational "
              "and sequential logic using HDL; writing RTL code for basic digital circuits.  "
              "Practical component: RTL Design and Implementation Labs (40 h), Design Synthesis "
              "and Optimization Labs (15 h), Timing Analysis and Closure Labs (10 h).",
              {"s": 9, "c": SLATE, "i": True})])

    w.callout("What's inside", [
        [B("Part 1  "), N("The Verilog language — modules, values, types, literals, vectors, "
                          "operators, width rules, procedural blocks, control flow, tasks and "
                          "functions, parameters, generate, directives, system tasks, and the "
                          "synthesisable subset")],
        [B("Part 2  "), N("Modelling logic — the inference map, combinational style, latch "
                          "inference, the clocked template, blocking vs non-blocking, the event "
                          "queue, reset strategy, counters, shift registers, edge detection, "
                          "CDC, state machines, memory and pipelining")],
        [B("Part 3  "), N("Writing RTL for basic circuits — the design catalogue, verification "
                          "method, testbench construction, scoreboards, randomisation, a "
                          "synchronous FIFO and a complete UART, plus two real bug case studies")],
        [B("Part 4  "), N("Tool setup and six guided tutorials — the open-source chain, Vivado "
                          "Design Suite and ModelSim, step by step at the keyboard")],
        [B("Part 5  "), N("60 practice exercises, graded from recall through prediction to "
                          "full design")],
        [B("Part 6  "), N("Full worked solutions to every exercise")],
        [B("Part 7  "), N("Reference — glossary, command card, coding-standard checklist and "
                          "the troubleshooting table")],
    ], color=TEAL)

    w.callout("How to use this workbook", [
        [N("Read a section, then immediately open the matching file in "),
         M("Topic4_Lab/rtl/"), N(" and read the real code. Every idea here has a file behind "
                                 "it. When the workbook quotes a tool output, run the command "
                                 "yourself and compare — the outputs quoted from Icarus "
                                 "Verilog 12.0, Verilator 5.020 and Yosys 0.33 were really "
                                 "produced by those tools on this code.")],
        [N("Work the exercises in order within each block. The prediction exercises "
           "(“what does this become?”) are the ones that build the skill this module is "
           "actually assessing; do not skip them because they have no code to type.")],
    ], color=AMBER, fill="FFF7EC", bar="C77514")
    w.page_break()

    # ============================================================ PART 1
    w.h1("Part 1 · The Verilog Language")
    w.para([N("Verilog exists to describe hardware. That sentence sounds obvious and is the "
              "single hardest thing for a newcomer with programming experience to internalise, "
              "because the language deliberately borrows C's syntax. This part covers the "
              "language systematically. Read it once end to end, then use it as a reference.",
              {"s": 10.5})])

    # ---------------------------------------------------------- 1.1
    w.h2("1.1  Why an HDL exists at all")
    w.para("Before hardware description languages, a chip was captured as a schematic: a drawing "
           "of every gate and every wire. That works up to a few thousand gates. A modern SoC has "
           "billions of transistors. Nobody draws that.")
    w.para("An HDL lets you describe behaviour and structure in text, which means the description "
           "can be version-controlled, diffed, reviewed, parameterised, generated by scripts and — "
           "crucially — automatically translated into gates by a synthesis tool. That last step is "
           "what makes the whole industry possible: you describe what the circuit does between "
           "clock edges, and a tool works out which gates in which technology will do it.")
    w.para("Verilog and VHDL are the two established HDLs. Verilog dominates in industry in Asia "
           "and North America and is what this course uses; SystemVerilog is its modern superset "
           "and adds features you will meet in verification. VHDL is more verbose, more strongly "
           "typed, and common in European and defence work. The concepts transfer completely; only "
           "the syntax differs.")

    w.h3("Verilog is not a programming language")
    w.table(["A program", "A hardware description"],
            [["Statements execute one after another",
              "Everything described exists at once and runs continuously"],
             ["A variable holds a value in memory",
              "A signal is a wire or a register that physically exists"],
             ["A function call runs code that already exists",
              "A function call creates ANOTHER COPY of the logic"],
             ["A loop repeats over time",
              "A loop is unrolled at compile time into repeated hardware"],
             ["Speed depends on how many instructions run",
              "Speed depends on the longest path between two registers"],
             ["Adding a feature costs memory",
              "Adding a feature costs area, power and possibly timing"]],
            widths=[3.3, 3.3], size=9.5, align_center=False)
    w.callout("The habit that prevents most beginner problems",
              ["Before you type a line of Verilog, sketch the hardware: which registers exist, "
               "what logic sits between them, what feeds what. Then write the code that describes "
               "that sketch. Designers who sketch first produce clean RTL immediately. Designers "
               "who type first and hope produce latches, races and timing failures — and then "
               "spend the afternoon in a waveform viewer."],
              color=GREEN, fill="EEF7F1", bar="2A9D5C")

    # ---------------------------------------------------------- 1.2
    w.h2("1.2  Levels of abstraction")
    w.image("abstraction_levels", 6.4,
            "The four levels of description. Only the middle two are what you write.")
    w.para("All four levels below are legal Verilog and all four simulate. Only some synthesise.")
    w.table(["Level", "You describe", "Synthesises?", "Who writes it"],
            [["Behavioural", "Algorithms, with delays and file I/O",
              "Mostly no", "Verification engineers, in testbenches"],
             ["RTL (register transfer)", "What happens to data between clock edges",
              "Yes — this is the target", "Design engineers. You."],
             ["Gate level", "Explicit AND, OR, NOT primitives and their connections",
              "Yes, but nobody writes it", "Synthesis tools, as output"],
             ["Switch level", "Individual transistors",
              "No", "Nobody, outside custom cell design"]],
            widths=[1.5, 2.3, 1.3, 2.0], size=9.5, align_center=False)
    w.para([N("RTL is the sweet spot: abstract enough to read and to port between technologies, "
              "concrete enough that the hardware is predictable. You should be able to "),
            B("read"), N(" gate-level netlists — you will look at them in synthesis reports — but "
                         "you will never write one.")])

    # ---------------------------------------------------------- 1.3
    w.h2("1.3  The module — the only unit of design")
    w.image("module_anatomy", 6.4, "Every design you write has this shape.")
    w.para("A module is a box with a name, a set of ports and a body. Everything you build is a "
           "box, or a box made of boxes. There is no other unit of design in Verilog.")
    w.code([
        "`timescale 1ns / 1ps            // time unit / precision -- first line of every file",
        "`default_nettype none           // no implicit wires -- see 1.14",
        "",
        "module adder4 #(",
        "    parameter integer W = 4     // parameters: sized at instantiation",
        ")(",
        "    input  wire [W-1:0] a,      // port list: the module's interface",
        "    input  wire [W-1:0] b,",
        "    input  wire         cin,",
        "    output wire [W-1:0] sum,",
        "    output wire         cout",
        ");",
        "",
        "    wire [W:0] tmp;             // internal signals: not visible outside",
        "",
        "    assign tmp  = a + b + cin;  // the design body: what it actually does",
        "    assign sum  = tmp[W-1:0];",
        "    assign cout = tmp[W];",
        "",
        "endmodule",
        "",
        "`default_nettype wire           // restore the default for other files",
    ], caption="A complete, well-formed Verilog module")

    w.h3("Port directions")
    w.table(["Direction", "Meaning", "Rules"],
            [["input", "Driven from outside, read inside",
              "Never assign to an input inside the module"],
             ["output", "Driven inside, read outside",
              "Exactly ONE thing inside may drive it"],
             ["inout", "Bidirectional", "Only on real chip pins; needs a tri-state driver"]],
            widths=[1.1, 2.6, 3.2], size=9.5, align_center=False)

    w.h3("ANSI and non-ANSI headers")
    w.para("Two styles exist. Use the ANSI style shown above, where direction, type and width all "
           "appear once in the header. You will meet the older non-ANSI style in legacy code:")
    w.code([
        "module adder4 (a, b, cin, sum, cout);      // non-ANSI: bare names in the header",
        "  input  [3:0] a, b;                       // ... then declared again in the body",
        "  input        cin;",
        "  output [3:0] sum;",
        "  output       cout;",
        "  ...",
        "endmodule",
    ], caption="Non-ANSI style — read it, do not write it")
    w.para("The non-ANSI form states each port twice, which means the two statements can disagree. "
           "Every modern coding standard requires ANSI headers.")

    # ---------------------------------------------------------- 1.4
    w.h2("1.4  Four-value logic")
    w.image("four_value_logic", 6.0, "Two hardware values, and two things the simulator says.")
    w.table(["Value", "Means", "Real hardware?"],
            [["0", "Logic low", "Yes"],
             ["1", "Logic high", "Yes"],
             ["x", "Unknown — the simulator cannot say what this is",
              "No. Real silicon always has SOME voltage."],
             ["z", "High impedance — nothing is driving this wire",
              "Yes, on a tri-state bus or an unconnected pin"]],
            widths=[0.8, 3.6, 2.6], size=9.5, align_center=False)
    w.para([B("x is a simulation concept, not a hardware state."), N(" Real silicon settles on "
            "some voltage; the simulator says x when it cannot determine which. An x on your "
            "waveform is the simulator telling you that YOU have not determined it either — that "
            "the design's behaviour depends on something undefined.")])
    w.h3("Where an x comes from — the five usual suspects")
    w.numbered([
        "A register that was never reset. Its power-up value is genuinely unknown, so the "
        "simulator says so.",
        "A wire that nothing drives — often a typo in a port connection, which "
        "`default_nettype none would have caught.",
        "A wire that TWO things drive, with different values. A multi-driver conflict resolves "
        "to x.",
        "Reading past the end of a vector or outside the bounds of an array.",
        "Arithmetic or logic on a value that is already x. x is contagious.",
    ])
    w.callout("Chasing an x", [
        [N("Because x spreads, the signal where you NOTICED the x is almost never the signal that "
           "CAUSED it. In the waveform viewer, find the first time the suspect signal goes x, then "
           "add whatever drives it and repeat, walking backwards, until you reach the first signal "
           "that turned x on its own. That signal is the bug.")],
        [N("In a testbench, always compare with "), M("==="), N(" and "), M("!=="),
         N(" — the four-state comparison operators. Plain "), M("=="),
         N(" against an x returns x, which is not true, so the check silently passes and the bug "
           "escapes.")],
    ], color=RED, fill="FDECEF", bar="C01F43")

    # ---------------------------------------------------------- 1.5
    w.h2("1.5  Nets and variables — wire, reg and the great naming accident")
    w.image("nets_vs_variables", 6.0, "Whether it becomes a flip-flop depends on HOW you assign.")
    w.para([B("A reg is not a register."), N(" It is a variable in the simulator — a thing that "
            "holds a value between assignments. Whether it becomes a flip-flop, a piece of "
            "combinational logic, or a latch depends entirely on how and where you assign it. The "
            "keyword was named in 1985 and the name has confused students ever since.")])
    w.table(["You want", "Declare it", "Assign it with", "You get"],
            [["A wire between two things", "wire", "assign, or a port connection", "A wire"],
             ["Combinational logic in a block", "reg", "= inside always @(*)", "Gates"],
             ["A flip-flop", "reg", "<= inside always @(posedge clk)", "A flip-flop"],
             ["An output driven by assign", "wire (the default)", "assign", "A wire"],
             ["An output driven by a block", "reg", "inside the always block", "Depends on block"]],
            widths=[1.9, 1.3, 2.2, 1.2], size=9, align_center=False)
    w.h3("Net types you will actually meet")
    w.table(["Type", "Use"],
            [["wire", "The default net. A simple connection with one driver."],
             ["tri", "Identical to wire; the name documents that several drivers are expected."],
             ["supply0 / supply1", "A constant tie to ground or to the supply rail."],
             ["wand / wor", "Wired-AND / wired-OR resolution. Rare in synthesisable RTL."]],
            widths=[1.4, 5.0], size=9.5, align_center=False)
    w.h3("Variable types")
    w.table(["Type", "Use"],
            [["reg", "The general variable. Any width. Synthesisable."],
             ["integer", "A 32-bit signed variable. Loop counters. Not for design signals."],
             ["real", "Floating point. Simulation only — never synthesises."],
             ["time", "A 64-bit unsigned value for timestamps. Simulation only."]],
            widths=[1.4, 5.0], size=9.5, align_center=False)
    w.callout("SystemVerilog fixes this",
              [[N("SystemVerilog adds "), M("logic"), N(", which may be driven either by an "
                  "assign or from inside a procedural block, and removes the wire/reg decision "
                  "entirely. If your tool flow accepts SystemVerilog, use "), M("logic"),
                N(" everywhere. You must still be able to READ Verilog-2001, because a very large "
                  "amount of existing IP is written in it — including much of what you will be "
                  "asked to maintain.")]],
              color=TEAL)

    # ---------------------------------------------------------- 1.6
    w.h2("1.6  Number literals")
    w.image("literals", 6.2, "Every constant has a width as well as a value.")
    w.para([N("The general form is "), M("<width>'<radix><value>"), N(". The radix is one of "),
            M("b"), N(" binary, "), M("o"), N(" octal, "), M("d"), N(" decimal, "), M("h"),
            N(" hexadecimal. Underscores may appear anywhere in the value and are ignored — use "
              "them to group digits.")])
    w.table(["Literal", "Width", "Value", "Note"],
            [["8'd10", "8", "00001010", "Decimal 10 in 8 bits"],
             ["8'hFF", "8", "11111111", "255"],
             ["4'b1010", "4", "1010", "10"],
             ["8'b1010_1100", "8", "10101100", "Underscores are cosmetic"],
             ["12", "32, SIGNED", "…00001100", "No width given — this is the trap"],
             ["-8'd1", "8", "11111111", "Two's complement of 1"],
             ["8'sd200", "8, signed", "11001000", "Signed: this is −56, not 200"],
             ["8'bx", "8", "xxxxxxxx", "x extends to fill the width"],
             ["'0  '1  'x  'z", "context", "all bits", "SystemVerilog fill literals"]],
            widths=[1.4, 1.1, 1.6, 2.3], size=9, align_center=False)
    w.callout("The house rule: size every literal",
              [[N("Write "), M("8'd0"), N(" not "), M("0"), N(". Write "), M("{W{1'b0}}"),
                N(" for a parameterised all-zeros. It costs three keystrokes and eliminates an "
                  "entire family of silent bugs — because an unsized literal is 32 bits wide and "
                  "signed, and both of those facts will eventually surprise you.")]],
              color=RED, fill="FDECEF", bar="C01F43")

    # ---------------------------------------------------------- 1.7
    w.h2("1.7  Vectors, slices, concatenation and replication")
    w.image("vector_ops", 6.2, "Taking buses apart and putting them back together.")
    w.code([
        "wire [7:0] d;                 // 8 bits, d[7] is the MSB   (use this order)",
        "wire [0:7] e;                 // 8 bits, e[0] is the MSB   (legal, confusing, avoid)",
        "",
        "d[3]                          // one bit",
        "d[7:4]                        // the top nibble -- a 4-bit value",
        "{a, b}                        // concatenation: a then b, side by side",
        "{4{2'b10}}                    // replication: 8'b10101010",
        "{2'b11, {3{1'b0}}, 3'b101}    // 8'b11000101 -- mix freely",
        "",
        "d[i +: 4]                     // indexed part-select: 4 bits UP from i",
        "d[i -: 4]                     // 4 bits DOWN from i",
    ], caption="Vector operations")
    w.h3("Arrays are not vectors")
    w.code([
        "reg [7:0] v;                  // ONE 8-bit vector -- you can slice it",
        "reg       v [0:7];            // EIGHT 1-bit elements -- you index it, cannot slice",
        "reg [7:0] m [0:255];          // 256 bytes -- a memory",
        "",
        "v[3:0]                        // legal: a slice of a vector",
        "m[17]                         // legal: one byte of the memory",
        "m[17][3:0]                    // legal: the bottom nibble of that byte",
        "m[3:0]                        // ILLEGAL: you cannot slice an array",
    ], caption="Vectors vs arrays")
    w.callout("Why indexed part-select exists",
              [[N("A plain slice needs constant bounds: "), M("d[i:j]"),
                N(" is illegal when i is a signal, because the tool must know how many bits the "
                  "expression produces. "), M("d[i +: 4]"),
                N(" has a constant WIDTH (4) and a variable POSITION — which is exactly what a "
                  "multiplexer does, so it synthesises to one. This is the correct idiom for "
                  "'select the nth byte of a wide bus'.")]],
              color=GREEN, fill="EEF7F1", bar="2A9D5C")

    # ---------------------------------------------------------- 1.8
    w.h2("1.8  Operators")
    w.image("operator_map", 6.4, "The operators you will actually use.")
    w.table(["Class", "Operators", "Result width"],
            [["Arithmetic", "+  -  *  /  %  **", "Widest operand (see 1.9)"],
             ["Bitwise", "~  &  |  ^  ~^", "Widest operand"],
             ["Logical", "!  &&  ||", "1 bit"],
             ["Relational", "<  <=  >  >=", "1 bit"],
             ["Equality", "==  !=  ===  !==", "1 bit"],
             ["Reduction", "&  ~&  |  ~|  ^  ~^  (unary)", "1 bit"],
             ["Shift", "<<  >>  <<<  >>>", "Left operand"],
             ["Conditional", "? :", "Widest branch"],
             ["Concatenation", "{ }  { { } }", "Sum of parts"]],
            widths=[1.3, 2.7, 2.3], size=9.5, align_center=False)
    w.h3("Logical versus bitwise — the classic slip")
    w.code([
        "a = 4'b0011;  b = 4'b1100;",
        "",
        "a && b   ->  1'b1        // both are non-zero, so both are 'true'",
        "a &  b   ->  4'b0000     // bit by bit: 0&1, 0&1, 1&0, 1&0",
        "",
        "!a       ->  1'b0        // a is non-zero, so 'not a' is false",
        "~a       ->  4'b1100     // every bit inverted",
    ], caption="&& gives one bit; & gives a vector")
    w.h3("Reduction operators are underused")
    w.table(["Expression", "Means", "Hardware"],
            [["|req", "Is ANY bit set?", "An OR tree — a 'valid' flag"],
             ["&full", "Are ALL bits set?", "An AND tree"],
             ["^data", "Parity of the whole vector", "An XOR tree"],
             ["~|result", "Is the result zero?", "A NOR tree — the ALU's Z flag"]],
            widths=[1.3, 2.3, 2.8], size=9.5, align_center=False)
    w.h3("Equality: == versus ===")
    w.para([M("=="), N(" is the 2-state comparison: if either operand contains an x or a z the "
                       "result is x. "), M("==="),
            N(" is the 4-state comparison: it compares x against x and z against z and returns a "
              "definite 1 or 0. Use "), M("=="),
            N(" in design code (where x should not exist and you want it to propagate if it does) "
              "and "), M("==="), N(" in testbenches (where you must be able to detect an x).")])
    w.h3("Signed arithmetic")
    w.para([N("Verilog is unsigned by default. A "), M("reg [7:0]"),
            N(" holding 8'hFF is 255, not −1. To get signed behaviour, declare "),
            M("reg signed [7:0]"), N(" or cast at the point of use with "), M("$signed(a)"),
            N(". Mixing signed and unsigned in one expression makes the WHOLE expression unsigned "
              "— which is why the ALU's set-on-less-than uses an explicit "),
            M("$signed(a) < $signed(b)"), N(".")])

    # ---------------------------------------------------------- 1.9
    w.h2("1.9  Width and sign rules — where silent bugs come from")
    w.image("width_rules", 6.0, "Verilog truncates silently, and never warns you.")
    w.para("Verilog evaluates an expression at a width determined by its context, then assigns "
           "that result to the target. If the target is narrower, the extra bits are discarded "
           "with no error, no warning and no runtime check. The simulation and the synthesised "
           "hardware agree with each other — and both are wrong.")
    w.numbered([
        "The width of a self-determined expression is the width of its widest operand.",
        "In an assignment, the right-hand side is evaluated at max(width of LHS, width of RHS) — "
        "so a narrow target can silently truncate.",
        "An unsized literal is 32 bits, signed.",
        "If ANY operand of an expression is unsigned, the whole expression is unsigned.",
        "Shifts, concatenations and comparisons all follow their own rules; when in doubt, be "
        "explicit.",
    ])
    w.code([
        "wire [3:0] a = 4'd9, b = 4'd8;",
        "wire [3:0] sum1 = a + b;         // 17 -> truncated to 4'd1. No warning.",
        "wire [4:0] sum2 = a + b;         // 17 -> 5'd17. Correct.",
        "wire [4:0] sum3 = {1'b0,a} + {1'b0,b};   // explicit, and always correct",
    ], caption="The truncation that catches everybody")
    w.callout("A real bug from this lab", [
        [N("The Topic 4 UART originally wrote its bit-timing limit as "),
         M("CLKS_PER_BIT[CW-1:0]"), N(". With "), M("CLKS_PER_BIT = 16"), N(", "), M("CW"),
         N(" is $clog2(16) = 4 — and 16 does not fit in 4 bits, so the expression truncated to "),
         B("zero"), N(".")],
        [N("HALF_BIT became 0, the receiver never waited half a bit before sampling, and it began "
           "sampling on bit boundaries instead of bit centres. 0x00 came back correctly; 0xFF "
           "came back as 0xF7. An intermittent, data-dependent failure — the worst kind.")],
        [N("The fix is to compute the limits as integers and slice them where they are USED: "),
         M("localparam integer FULL_BIT = CLKS_PER_BIT - 1;"), N(" and then "),
         M("clk_cnt == FULL_BIT[CW-1:0]"), N(". See "), M("rtl/uart_rx.v"), N(".")],
        [B("Verilator would have reported WIDTHTRUNC on that line in one second. Lint first.")],
    ], color=RED, fill="FDECEF", bar="C01F43")
