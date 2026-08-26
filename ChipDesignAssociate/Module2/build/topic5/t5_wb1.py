# -*- coding: utf-8 -*-
"""Topic 5 workbook — front matter + Part 1: verification fundamentals."""
import _boot
from wbkit import *
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL


def B(t, d=None, **kw):
    kw.update(d or {}); kw["b"] = True; return (t, kw)


def N(t, d=None, **kw):
    kw.update(d or {}); return (t, kw)


def I(t, d=None, **kw):
    kw.update(d or {}); kw["i"] = True; return (t, kw)


def M(t, d=None, **kw):
    kw.update(d or {}); kw["f"] = MONOF; return (t, kw)


def build(w):
    # ------------------------------------------------------------ cover
    w.para([N("CHIP DESIGN ASSOCIATE  ·  O-LEVEL ‘CHIP DESIGN’",
              {"b": True, "s": 11, "c": TEAL})], space_after=2)
    p = w.d.add_paragraph()
    r = p.add_run("Module 2 — Topic 5")
    r.font.name = HEADF; r.font.size = Pt(15); r.font.bold = True; r.font.color.rgb = SLATE
    p.paragraph_format.space_after = Pt(2)
    p = w.d.add_paragraph()
    r = p.add_run("RTL Simulation and Verification")
    r.font.name = HEADF; r.font.size = Pt(25); r.font.bold = True; r.font.color.rgb = NAVY
    p.paragraph_format.space_after = Pt(4)
    p = w.d.add_paragraph()
    r = p.add_run("Tutorial & Practice Workbook")
    r.font.name = HEADF; r.font.size = Pt(16); r.font.color.rgb = AMBER
    p.paragraph_format.space_after = Pt(10)
    w.para([N("A self-study companion to the Topic 5 slide deck. It explains every verification "
              "technique the deck introduces and why it exists, walks through seven guided "
              "tutorials at the keyboard, and ends with 60 graded exercises and full worked "
              "solutions. Every result quoted here was produced by running the code in "
              "Topic5_Lab/. Nothing in this workbook requires you to look anything up "
              "elsewhere.", {"s": 10.5})])
    w.para([N("NOS: NIE/ELE/N0102  ·  Module 2 “Verilog RTL coding for Synthesis”, Topic 5 "
              "“RTL Simulation and Verification”  ·  Syllabus: functional verification "
              "techniques for RTL designs; introduction to test benches and testbench "
              "development; simulation and debugging of RTL designs.  Practical component: "
              "validate RTL designs through simulation using test benches and debugging "
              "techniques; develop and execute test benches for comprehensive simulation and "
              "debugging of Verilog RTL code.",
              {"s": 9, "c": SLATE, "i": True})])

    w.callout("What's inside", [
        [B("Part 1  "), N("Verification fundamentals — why it dominates the schedule, the "
                          "verification gap, the loop, directed versus random, the four ways to "
                          "check an answer, coverage, the verification plan, and metrics")],
        [B("Part 2  "), N("Testbench development — the six parts, clock and reset, stimulus "
                          "timing, reference models and scoreboards, constrained-random and "
                          "seeds, functional coverage, assertions, layered environments, "
                          "SystemVerilog, UVM and formal")],
        [B("Part 3  "), N("Simulation and debugging — the event engine, the stratified queue, "
                          "waveform formats and viewers, a debugging procedure, chasing an x, "
                          "symptom-to-cause, and simulation performance")],
        [B("Part 4  "), N("Tool setup and seven guided tutorials — the open-source chain, "
                          "Vivado xsim and ModelSim, step by step at the keyboard")],
        [B("Part 5  "), N("60 practice exercises, graded from recall through diagnosis to full "
                          "environment construction")],
        [B("Part 6  "), N("Full worked solutions to every exercise")],
        [B("Part 7  "), N("Reference — glossary, command card, testbench checklist and the "
                          "troubleshooting table")],
    ], color=TEAL)

    w.callout("The one idea this whole topic rests on", [
        [B("A testbench is not finished when it passes. It is finished when it would FAIL if "
           "the design were wrong.")],
        [N("Topic5_Lab contains one correct FIFO and five copies with a single realistic bug "
           "each. All five lint clean, all five synthesise, and all five pass a testbench most "
           "students would call finished. Run "), M("./scripts/clinic.sh"),
         N(" before you read any further — the table it prints is the argument the rest of "
           "this workbook exists to explain.")],
    ], color=RED, fill="FDECEF", bar="C01F43")
    w.page_break()

    # ============================================================ PART 1
    w.h1("Part 1 · Verification Fundamentals")
    w.para([N("Topic 4 taught you to write RTL. This topic asks the harder question: "),
            B("how do you know it works?"),
            N("  The answer is not \"the simulation passed\", and Part 1 is about why not.",
              {"s": 10.5})])

    # ---------------------------------------------------------- 1.1
    w.h2("1.1  Why verification is half the job")
    w.image("why_verify", 6.4, "Where the effort goes, and what a bug costs by the stage that "
                               "finds it.")
    w.para("On a typical ASIC or FPGA block, verification consumes more engineering effort than "
           "design. That is not a sign of inefficiency; it is what the economics demand. A bug "
           "is the same bug at every stage of the flow. What changes is what it costs to find "
           "and fix.")
    w.table(["Found by", "What it costs you", "Relative"],
            [["A linter", "one second, with a file name and a line number", "×1"],
             ["Simulation", "a minute, plus the time to read the waveform", "×3"],
             ["Synthesis", "a longer run, and a structural surprise to understand", "×8"],
             ["The FPGA lab", "re-synthesise, re-place, re-programme, re-instrument", "×25"],
             ["Silicon", "a mask respin: months, and a great deal of money", "×60"]],
            widths=[1.7, 3.6, 1.1], size=9.5, align_center=False)
    w.callout("The consequence for how you work",
              ["Run the cheap checks first, always, and do not skip one because you are in a "
               "hurry. A bug that a one-second lint would have named should never reach a "
               "waveform viewer, and a bug a testbench would have caught should never reach a "
               "board."],
              color=AMBER, fill="FFF7EC", bar="C77514")

    # ---------------------------------------------------------- 1.2
    w.h2("1.2  The verification gap")
    w.image("verification_gap", 6.4, "The problem, and what is done about it.")
    w.para("Exhaustive testing is impossible for anything real, and the arithmetic is not close. "
           "A single 32-bit adder has 2^64 input pairs; at a billion per second, testing all of "
           "them takes over 500 000 years. A design with 100 flip-flops has more reachable "
           "states than there are atoms in the observable universe. Adding one flip-flop doubles "
           "the state space; adding one engineer does not.")
    w.para("So verification is not proof. It is measured evidence, gathered by four means:")
    w.numbered([
        "DIRECTED tests, for the cases you can name — boundaries, reset, protocol corners.",
        "CONSTRAINED-RANDOM stimulus, for the cases you cannot name, generated within "
        "constraints you do specify.",
        "ASSERTIONS, which state the rules so they are checked on every clock edge of every "
        "test rather than at the points you remembered to look.",
        "COVERAGE, which measures what the tests actually reached — and is the only one of the "
        "four that can tell you when to stop.",
    ])
    w.callout("Verification and validation are different questions",
              [[B("Verification"), N(" asks: does the design do what the specification says? "),
                B("Validation"), N(" asks: was the specification the right one? This topic is "
                                   "entirely about the first. The second is usually somebody "
                                   "else's job, and a design can pass verification perfectly "
                                   "while solving the wrong problem.")]],
              color=TEAL)

    # ---------------------------------------------------------- 1.3
    w.h2("1.3  The verification loop")
    w.image("ver_flow", 6.4, "It ends at coverage, not at PASS.")
    w.para("Notice where the loop ends. Not at \"the simulation passed\" — at \"coverage is "
           "closed\". A pass with unmeasured coverage tells you only that nothing broke in "
           "whatever cases you happened to run, which may be very few of the interesting ones.")
    w.para("Notice also the feedback arrow. Coverage does not merely report; it tells you what "
           "to do next. Every MISS is a specific piece of stimulus somebody has to write. That "
           "is what makes it a loop rather than a final report, and it is why the plan comes "
           "first.")

    # ---------------------------------------------------------- 1.4
    w.h2("1.4  Directed and constrained-random")
    w.image("directed_vs_random", 6.4, "Precise darts, and a cloud that reaches further.")
    w.table(["", "Directed", "Constrained-random"],
            [["You write", "the case AND the expected answer", "the constraints; a model gives "
              "the answer"],
             ["Cost per test", "high — one test, one case", "low — one test, thousands of cases"],
             ["Reproducible", "always", "yes, from the seed"],
             ["Finds", "what you thought of", "what you did not"],
             ["Debugging a failure", "easy — you know what you asked for",
              "harder — but the seed replays it exactly"],
             ["Use it for", "boundaries, reset, protocol corners, regressions of known bugs",
              "volume, and the long tail of combinations"]],
            widths=[1.5, 2.4, 2.5], size=9.5, align_center=False)
    w.para("A real project uses both, in that order: directed tests first, because they are "
           "cheap and they fail fast on the obvious things; then random, to cover the "
           "combinations nobody enumerated. In Topic5_Lab, Lab V2 is the directed stage and Lab "
           "V3 the random one, and the measured difference between them is one bug out of five.")

    # ---------------------------------------------------------- 1.5
    w.h2("1.5  Four ways to decide whether an answer was right")
    w.image("checker_taxonomy", 6.4, "From weakest to strongest.")
    w.para("This is the single biggest determinant of how good a testbench is — not how many "
           "cases it runs, but how it decides whether each one was correct.")
    w.table(["Method", "Scales to", "Fails when"],
            [["A human reads the waveform", "one run, once",
              "always: it cannot run overnight or in a regression"],
             ["Hard-coded expected values", "about twenty cases",
              "the values themselves acquire bugs, silently"],
             ["A reference model + scoreboard", "millions of cycles",
              "the model was written by reading the RTL"],
             ["Assertions", "every cycle of every test, for ever",
              "the rule you did not think to state"]],
            widths=[2.3, 1.8, 2.6], size=9.5, align_center=False)
    w.callout("The line that separates a testbench from a demonstration",
              ["A testbench must print PASS or FAIL, and its verdict must be a line of text a "
               "machine can grep for. If deciding whether the run was correct requires a human "
               "to look at a waveform, you cannot put it in a regression, you cannot run it "
               "overnight, and you will stop running it within a week."],
              color=RED, fill="FDECEF", bar="C01F43")

    # ---------------------------------------------------------- 1.6
    w.h2("1.6  Coverage")
    w.image("coverage_types", 6.4, "Only one of the two knows the specification.")
    w.h3("Code coverage — computed for you")
    w.table(["Kind", "Asks"],
            [["line", "was this line ever executed?"],
             ["branch", "was each arm of each if/else and case taken?"],
             ["condition (expression)", "did each sub-expression take both values?"],
             ["toggle", "did each bit go 0→1 and 1→0?"],
             ["FSM state / transition", "was each state entered, and each arc taken?"]],
            widths=[2.0, 4.4], size=9.5, align_center=False)
    w.para("Code coverage is free — the tool computes it. It is also weak on its own, and the "
           "reason is worth stating precisely: it measures the code you WROTE. It cannot measure "
           "behaviour you forgot to write code for, and it says nothing about whether the "
           "outputs were checked. A testbench that runs every line and checks nothing reports "
           "100%.")
    w.h3("Functional coverage — written by you")
    w.para("Functional coverage is a list of specification-level situations you decided in "
           "advance ought to happen, each with a counter. Nobody but you can write it, because "
           "it comes from the specification, not from the code.")
    w.bullets([
        "Did the FIFO ever reach full? Ever reach empty?",
        "Did a read and a write ever occur on the same cycle? ...while empty? ...while full?",
        "Did the pointers ever wrap?",
        "Did we ever go full → empty → full? (a sequence, not a state)",
        "Was reset ever asserted in the middle of a transfer?",
    ])
    w.callout("Coverage holes are not failures",
              ["A hole means \"this never happened during testing\". That is a gap in your "
               "evidence, not a bug in the design. You close it by writing stimulus that reaches "
               "it, or by waiving it in writing with a reason — for example, because the "
               "situation is impossible by construction. What you must not do is leave it "
               "unexplained."],
              color=GREEN, fill="EEF7F1", bar="2A9D5C")

    # ---------------------------------------------------------- 1.7
    w.h2("1.7  The verification plan")
    w.image("ver_plan", 6.4, "One table, written before the testbench.")
    w.para("A verification plan does not have to be a document. It has to be a table with one "
           "row per feature in the specification, and columns saying how each feature is "
           "checked, what stimulus reaches it, which coverage bin proves it happened, and "
           "whether that is done yet.")
    w.para([B("Write it before the testbench."), N("  A plan written afterwards documents what "
            "you happened to build. A plan written first tells you what to build, and the gaps "
            "are visible while they are still cheap to close.")])
    w.para([B("Have somebody else review it."), N("  The author of a plan cannot see what the "
            "plan omits — that is the definition of an omission. This is the single cheapest "
            "quality measure available to a small team.")])

    # ---------------------------------------------------------- 1.8
    w.h2("1.8  The stages, and what each one catches")
    w.image("bug_escape", 6.4, "Each stage is a filter with a different mesh.")
    w.table(["Stage", "Catches", "Cannot catch"],
            [["Lint", "width truncation, inferred latches, undriven and unused signals, "
                      "incomplete sensitivity lists",
              "anything about behaviour — it never runs the design"],
             ["Simulation", "wrong behaviour on the stimulus you supplied",
              "behaviour you never stimulated; metastability"],
             ["Assertions", "a stated rule broken, at the cycle it breaks",
              "a rule you did not state"],
             ["Coverage", "situations that never occurred at all",
              "whether the check for that situation was correct"],
             ["Synthesis", "structure the simulator hid: latches, surprising flip-flop counts",
              "functional errors"],
             ["Formal", "a proof, or the shortest counter-example, for a stated property",
              "wide datapaths — it runs out of memory instead"]],
            widths=[1.2, 3.0, 2.4], size=9, align_center=False)

    # ---------------------------------------------------------- 1.9
    w.h2("1.9  The experiment this topic is built on")
    w.image("clinic_matrix", 6.4, "Measured, not asserted. Reproduce it with ./scripts/clinic.sh")
    w.para("Topic5_Lab contains one correct FIFO and five copies, each with a single realistic "
           "bug. Three testbenches of increasing strength are run against all six designs. The "
           "table is the result.")
    w.table(["Testbench", "Bugs caught", "What it added"],
            [["V1  naive directed", "0 of 5", "the six structural parts, and a verdict"],
             ["V2  model + corners", "4 of 5", "a reference model, and the boundary cases"],
             ["V3  constrained-random", "5 of 5", "weighted random stimulus, seeded"]],
            widths=[2.2, 1.4, 2.8], size=9.5, align_center=False)
    w.para("Every testbench passes on the golden design. That column matters as much as the "
           "others: a checker that fires on correct hardware is worse than no checker, because "
           "the team learns to ignore it.")
    w.h3("What changed at each step")
    w.table(["Step", "The decision", "Why it mattered"],
            [["V1 → V2", "expected values now come from an independent model of the spec, and "
                         "the test visits the boundaries",
              "you can check every cycle instead of the few you wrote checks for; and the bugs "
              "were all at the boundaries"],
             ["V2 → V3", "stimulus is weighted-random and seeded",
              "the fifth bug only shows when a read and a write are asserted together on an "
              "EMPTY FIFO — nobody writes that as a directed test"],
             ["V3 → V4", "a coverage model is sampled and reported",
              "a run that passes now also says what it reached; the write-heavy profile passes "
              "and never reaches empty at all"],
             ["V3 → V6", "assertions are bound to the DUT",
              "a broken rule is reported at the cycle it breaks, by name, instead of later at "
              "an output as a wrong number"]],
            widths=[1.0, 2.4, 3.2], size=9, align_center=False)

    # ---------------------------------------------------------- 1.10
    w.h2("1.10  How a project decides it is finished")
    w.para("\"We ran all the tests and they passed\" is not a sign-off criterion. These are, and "
           "every one of them is a number somebody outside the team can read and challenge.")
    w.table(["Metric", "Typical sign-off bar"],
            [["Functional coverage", "100% of the plan, or every hole waived in writing with a "
                                     "reason"],
             ["Code coverage", "95–100%, with the remainder explained line by line"],
             ["Assertion coverage", "every assertion evaluated at least once — an assertion that "
                                    "never fired proves nothing"],
             ["Regression pass rate", "100%, sustained over time, not \"usually green\""],
             ["Bug discovery rate", "flattened, and staying flat, over several weeks"],
             ["Bug severity trend", "no high-severity findings for an agreed period"]],
            widths=[2.2, 4.2], size=9.5, align_center=False)
    w.callout("The bug-rate curve is the one experienced teams watch",
              ["If you are still finding serious bugs at the same rate as last month, you are "
               "not near the end however good the coverage number looks — it means the coverage "
               "model is missing a category. If the rate has flattened AND coverage is closed "
               "AND the regression has been green for weeks, you have as much evidence as this "
               "method can produce."],
              color=AMBER, fill="FFF7EC", bar="C77514")
