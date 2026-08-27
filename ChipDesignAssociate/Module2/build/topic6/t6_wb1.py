# -*- coding: utf-8 -*-
"""Topic 6 workbook — front matter + Part 1: the physics of a flip-flop."""
import _boot
from wbkit import *
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL


def B(t, d=None, **kw):
    kw.update(d or {}); kw["b"] = True; return (t, kw)


def N(t, d=None, **kw):
    kw.update(d or {}); return (t, kw)


def I(t, d=None, **kw):
    kw.update(d or {}); kw["i"] = True; return (t, kw)


def M(t, d=None, **kw):
    kw.update(d or {}); kw["f"] = MONOF; return (t, kw)


def build(w):
    # ------------------------------------------------------------ cover
    w.para([N("CHIP DESIGN ASSOCIATE  ·  O-LEVEL ‘CHIP DESIGN’",
              {"b": True, "s": 11, "c": TEAL})], space_after=2)
    p = w.d.add_paragraph()
    r = p.add_run("Module 2 — Topic 6")
    r.font.name = HEADF; r.font.size = Pt(15); r.font.bold = True; r.font.color.rgb = SLATE
    p.paragraph_format.space_after = Pt(2)
    p = w.d.add_paragraph()
    r = p.add_run("Timing Constraints and Analysis")
    r.font.name = HEADF; r.font.size = Pt(25); r.font.bold = True; r.font.color.rgb = NAVY
    p.paragraph_format.space_after = Pt(4)
    p = w.d.add_paragraph()
    r = p.add_run("Tutorial & Practice Workbook")
    r.font.name = HEADF; r.font.size = Pt(16); r.font.color.rgb = AMBER
    p.paragraph_format.space_after = Pt(10)
    w.para([N("A self-study companion to the Topic 6 slide deck. It explains every timing "
              "concept the deck introduces and why it exists, walks you through building a "
              "working static timing analyser from nothing, and ends with 62 graded "
              "exercises and full worked solutions. Every number quoted in this workbook "
              "was produced by running the code in Topic6_Lab/. Nothing here requires you "
              "to look anything up elsewhere.", {"s": 10.5})])
    w.para([N("NOS: NIE/ELE/N0102  ·  Module 2 “Verilog RTL coding for Synthesis”, Topic 6 "
              "“Timing Constraints and Analysis”  ·  Syllabus: introduction to timing "
              "constraints in RTL design; timing analysis and optimization techniques; "
              "setup and hold time violations and resolution.  Practical component: Timing "
              "Analysis and Closure Labs (10 h) — timing analysis using industry-standard "
              "tools, constraint development and application, timing closure techniques; "
              "Design Synthesis and Optimisation Labs (15 h).",
              {"s": 9, "c": SLATE, "i": True})])

    w.callout("What's inside", [
        [B("Part 1  "), N("The physics — why a flip-flop has a sampling window, what setup "
                          "and hold really mean, where clock skew and jitter come from, and "
                          "the one equation the whole topic rests on")],
        [B("Part 2  "), N("Constraints — create_clock, uncertainty, input and output delay, "
                          "false paths, multicycle paths, and the order to write and debug "
                          "a constraint file")],
        [B("Part 3  "), N("Static timing analysis — the delay model, the timing graph, the "
                          "forward and backward sweeps, and how to read any timing report")],
        [B("Part 4  "), N("Optimisation — the seven fixes for a setup violation, in cost "
                          "order, with measured results for each")],
        [B("Part 5  "), N("Setup and hold violations — why they are opposite, why the clock "
                          "cannot fix hold, and a diagnosis procedure")],
        [B("Part 6  "), N("Tools — installation, the open-source flow, Vivado, OpenSTA, "
                          "and how SDC maps onto XDC")],
        [B("Part 7  "), N("Seven guided tutorials at the keyboard: you build the Liberty "
                          "file, the timing graph, the engine, then use it on real "
                          "violations")],
        [B("Exercises  "), N("62 graded exercises across seven parts, every one with a "
                             "worked solution and the exact command to check it")],
    ], color=NAVY, bar="0E2A47")

    w.callout("How to use this workbook", [
        [N("Read a part, then do its tutorial with the terminal open. The tutorials are "
           "written to be typed, not skimmed — every command is one you run, and every "
           "output shown is one you will see.")],
        [N("Exercises marked "), B("[H]"), N(" are hand calculations: do them on paper "
           "before you run anything. Exercises marked "), B("[C]"),
         N(" need the computer. Exercises marked "), B("[W]"),
         N(" ask you to write something down and defend it — these are the ones that turn "
           "a student into an engineer.")],
        [N("If a number you measure disagrees with a number printed here, do not assume the "
           "workbook is right. Find out which of you is wrong. That investigation is worth "
           "more than the answer.")],
    ], color=GREEN, fill="EEF7F1", bar="2A9D5C")

    w.page_break()

    # ================================================================ Part 1
    w.h1("Part 1 · What a Flip-Flop Actually Does")

    w.h2("1.1  The mental model you have, and the one you need")

    w.para([N("Until now you have thought of a flip-flop like this: at the rising edge of "
              "the clock, D is copied to Q. That model is good enough to write correct RTL "
              "and it is what a zero-delay simulator implements. It is also, physically, "
              "not true.")])
    w.para([N("A real flip-flop is a pair of cross-coupled latches. For it to settle into "
              "a definite state, the value on D must be "), B("still"),
            N(" for a short interval around the clock edge — before it and after it. "
              "Change D inside that interval and the flop does not cleanly capture either "
              "the old value or the new one. It may take an unbounded time to settle, or "
              "settle to the wrong value. This is called "), B("metastability"),
            N(", and it is what timing analysis exists to prevent.")])

    w.image("setup_hold_window", width=6.4,
            caption="The sampling window. Everything in this topic follows from this picture.")

    w.h3("The two halves of the window")
    w.table(["Name", "What it requires", "Symbol in a datasheet", "Typical value"],
            [["Setup time", "D stable BEFORE the clock edge", "t_su, t_setup",
              "0.02 – 0.2 ns"],
             ["Hold time", "D stable AFTER the clock edge", "t_h, t_hold",
              "0.00 – 0.1 ns, sometimes negative"],
             ["Clock-to-Q", "how long after the edge Q is valid", "t_co, t_cq",
              "0.05 – 0.3 ns"]],
            widths=[1.3, 2.6, 1.5, 1.4], size=9.5, bold_cols=(0,))

    w.callout("Hold time can be negative — and it is not a typo", [
        [N("Inside the cell, the clock takes a little time to reach the input latch. If "
           "that internal clock delay is longer than the latch's own requirement, D is "
           "allowed to change slightly BEFORE the external clock edge and still be "
           "captured correctly. The datasheet then quotes a negative hold time. "
           "A negative hold time is free margin: it makes hold violations harder to "
           "create.")],
    ], color=TEAL)

    w.h2("1.2  The two failures, and why they are not variations of one thing")

    w.para([N("Break the first half of the window and you have a "), B("setup violation"),
            N(": the data arrived too late, so the flop captured whatever was there before "
              "— the previous value. Break the second half and you have a "),
            B("hold violation"), N(": the data changed too soon after the edge, so the flop "
              "captured the NEW value, a whole clock cycle early.")])

    w.image("setup_vs_hold", width=6.4,
            caption="They look similar in a report and behave nothing alike in silicon.")

    w.callout("The difference that decides everything you do next", [
        [B("A setup violation means the chip is too slow. "),
         N("Run the clock slower and it works. You have shipped a product that is less "
           "competitive, but it is a working product.")],
        [B("A hold violation means the chip is broken. "),
         N("There is no clock frequency at which it works — not 1 MHz, not 1 Hz. The data "
           "races the clock edge and wins, and slowing the clock does not change the race. "
           "A hold violation that reaches silicon is a re-spin.")],
    ], color=RED, fill="FDECEF", bar="D6224A")

    w.h2("1.3  Where the time goes: the four parts of a path")

    w.para([N("Between one flip-flop and the next, four things consume time. Every timing "
              "number you will ever compute is a combination of these four.")])

    w.image("timing_path", width=6.5)

    w.numbered([
        [B("Clock-to-Q  "), N("(t_cq). The launching flop does not present its new output "
           "at the instant of the edge. It takes 0.145 ns in the library used throughout "
           "this workbook.")],
        [B("Logic delay.  "), N("Every gate between the two flops, plus the wires between "
           "them. In a 32-bit ripple-carry adder this is 32 gate delays in series — it "
           "dominates everything else.")],
        [B("Setup time  "), N("(t_su). The capturing flop needs the data early. This is "
           "subtracted from the time you have, so it behaves exactly like extra logic "
           "delay.")],
        [B("Clock skew.  "), N("The two flops do not see the same clock edge at the same "
           "instant, because the clock reached them through different amounts of wire and "
           "buffering.")],
    ])

    w.h2("1.4  The equation")

    w.para([N("Put those four together and you get the setup check. It is worth writing "
              "this out once in full, because every timing report you ever read is this "
              "equation with the numbers filled in.")])

    w.code([
        "arrival time    =  clock-to-Q  +  logic delay",
        "                   (when the data actually turns up at the capture flop's D pin)",
        "",
        "required time   =  period  +  skew  -  setup  -  uncertainty",
        "                   (the latest it could have turned up and still been captured)",
        "",
        "setup slack     =  required time  -  arrival time",
        "",
        "                =  ( period + skew - setup - uncertainty )",
        "                 - ( clock-to-Q + logic delay )"],
        caption="the setup check")

    w.para([N("And the hold check, which is deliberately laid out the same way so you can "
              "see what is missing from it:")])

    w.code([
        "arrival time    =  clock-to-Q  +  logic delay      <-- the SHORTEST path now",
        "",
        "required time   =  skew  +  hold  +  uncertainty",
        "",
        "hold slack      =  arrival time  -  required time",
        "",
        "#  Note two things:",
        "#    1. the clock period does not appear anywhere in the hold check",
        "#    2. the subtraction is the other way round"],
        caption="the hold check")

    w.callout("Read those two boxes until the difference is obvious", [
        [N("Setup wants the arrival time to be "), B("small"),
         N(". Hold wants it to be "), B("large"), N(". Every fix you apply to one makes "
           "the other worse. That tension is the whole of timing closure.")],
        [N("And because the period is in one equation and not the other, the clock "
           "frequency is a lever for setup and no lever at all for hold.")],
    ], color=NAVY, bar="0E2A47")

    w.h2("1.5  Slack, WNS and TNS")

    w.image("slack_equation", width=6.4)

    w.para([N("Slack is computed at every "), B("endpoint"),
            N(" in the design — every flip-flop's D pin and every output port. A design "
              "with a million flops produces a million slack numbers. Two summary numbers "
              "are extracted from that list:")])

    w.table(["Number", "Definition", "What it tells you"],
            [["WNS", "the worst (most negative) slack in the design",
              "how bad the worst single path is"],
             ["TNS", "the sum of every negative slack",
              "how many paths are bad, and by how much in total"],
             ["Fmax", "1 / (longest path delay)",
              "the highest clock frequency the design can run at"]],
            widths=[0.9, 2.9, 3.0], size=9.5, bold_cols=(0,))

    w.image("wns_tns", width=6.3)

    w.callout("Why you need both numbers", [
        [B("WNS −0.42, TNS −0.42.  "),
         N("One path is bad. Find it, fix it, and you are finished this afternoon.")],
        [B("WNS −0.42, TNS −180.  "),
         N("The same worst path, but hundreds of others are failing too. No local fix will "
           "help. Either the frequency target is wrong for this technology, or the "
           "architecture puts too much work in one cycle. This is a design review, not a "
           "debugging session.")],
    ], color=AMBER, fill="FFF7EC", bar="C77514")

    w.h2("1.6  Clock skew: the same clock, at different instants")

    w.para([N("A clock signal has to reach every flip-flop on the chip. It does so through "
              "a tree of buffers and a great deal of wire. Two flops at opposite corners "
              "of a die may see the same edge tens or hundreds of picoseconds apart. "
              "That difference is "), B("skew"), N(".")])

    w.image("clock_skew", width=6.4)

    w.para([N("Skew is signed, and its sign depends on which flop you are looking at. If "
              "the CAPTURE flop's clock arrives late relative to the launch flop's:")])
    w.bullets([
        [B("Setup gets easier. "), N("The capture edge happens later, so the data has more "
           "time to arrive. The skew adds to the required time.")],
        [B("Hold gets harder. "), N("The capture edge happens later, so the NEW data from "
           "the launch flop may already have arrived by the time the capture flop looks. "
           "The skew is subtracted from the hold margin.")],
    ])

    w.callout("Useful skew", [
        [N("Because skew helps setup, tools sometimes add it deliberately: delay the clock "
           "to a flop that is on a tight path, and steal margin from the next stage, which "
           "had some to spare. This is called useful skew or clock-tree scheduling. It is "
           "a real technique, and it is also the reason hold violations sometimes appear "
           "on paths you never touched.")],
    ], color=VIOLET, bar="7A4FBF")

    w.h2("1.7  Jitter and uncertainty")

    w.para([N("Skew is a fixed difference: the same on every cycle. "), B("Jitter"),
            N(" is a random one — the clock edge does not land in exactly the same place "
              "every cycle, because the oscillator and the PLL that produce it are not "
              "perfect. And before place-and-route runs, the clock tree does not exist at "
              "all, so its skew is not merely unknown — it is undefined.")])

    w.image("uncertainty", width=6.3)

    w.para([N("Both are handled by one constraint:")])
    w.code([
        "set_clock_uncertainty 0.150 -setup [get_clocks sys_clk]",
        "set_clock_uncertainty 0.050 -hold  [get_clocks sys_clk]"])

    w.table(["Design stage", "Typical setup uncertainty", "Why that number"],
            [["FPGA, post-synthesis", "0.10 – 0.20 ns",
              "no clock tree yet; jitter unknown"],
             ["FPGA, post-place-and-route", "0.02 – 0.05 ns",
              "the tool now models the real tree; only jitter remains"],
             ["ASIC, before clock-tree synthesis", "0.15 – 0.30 ns",
              "skew placeholder plus on-chip variation margin"],
             ["ASIC, sign-off", "jitter + OCV derating",
              "the tree is built and characterised"]],
            widths=[1.9, 1.9, 3.0], size=9.2, bold_cols=(0,))

    w.callout("The temptation, and why to resist it", [
        [N("When a design misses timing by 0.08 ns and the uncertainty is 0.15 ns, reducing "
           "the uncertainty to 0.05 makes the report pass instantly. It also makes the "
           "report a work of fiction. The margin was there to cover jitter and unmodelled "
           "skew; those did not go away because you edited a file.")],
        [N("Reduce uncertainty only when you can point at the measurement that justifies "
           "the smaller number.")],
    ], color=RED, fill="FDECEF", bar="D6224A")

    w.h2("1.8  PVT corners: the same netlist on different silicon")

    w.para([N("Two chips from the same wafer are not identical, and the same chip behaves "
              "differently at 0 °C and 125 °C, or at 0.9 V and 1.1 V. Timing is therefore "
              "not analysed once but at several "), B("corners"),
            N(" — combinations of process, voltage and temperature.")])

    w.image("corners", width=6.4)

    w.table(["Corner", "Process", "Voltage", "Temperature", "Signs off"],
            [["SLOW / worst", "slow", "minimum", "maximum", "SETUP"],
             ["FAST / best", "fast", "maximum", "minimum", "HOLD"],
             ["TYPICAL", "nominal", "nominal", "nominal", "nothing on its own"]],
            widths=[1.4, 1.1, 1.1, 1.4, 1.5], size=9.2, bold_cols=(0,))

    w.para([N("The logic is straightforward once you see it. Setup is a race between the "
              "data and the next clock edge, so the danger is the data being slow — check "
              "it where the logic is slowest. Hold is a race between the new data and the "
              "capture flop's hold requirement on the SAME edge, so the danger is the data "
              "being fast — check it where the logic is fastest. A design that passes both "
              "extremes passes everything in between.")])

    w.h2("1.9  Vocabulary you now own")

    w.table(["Term", "One-line definition"],
            [["setup time", "how long D must be stable before the clock edge"],
             ["hold time", "how long D must stay stable after the clock edge"],
             ["clock-to-Q", "delay from the clock edge to a valid Q"],
             ["arrival time", "when data actually reaches a pin, from the clock edge"],
             ["required time", "the latest (setup) or earliest (hold) it may arrive"],
             ["slack", "required − arrival for setup; arrival − required for hold"],
             ["WNS", "the worst negative slack in the design"],
             ["TNS", "the sum of all negative slacks"],
             ["Fmax", "1 / longest path delay"],
             ["critical path", "the path with the worst slack"],
             ["clock skew", "fixed difference in clock arrival between two registers"],
             ["clock jitter", "random cycle-to-cycle variation in the clock edge"],
             ["uncertainty", "margin covering jitter and unmodelled skew"],
             ["metastability", "the undefined state a flop enters if the window is broken"],
             ["PVT corner", "a process/voltage/temperature combination to analyse at"]],
            widths=[1.6, 5.2], size=9.2, bold_cols=(0,), align_center=False)

    w.callout("Part 1 self-check — answer these before moving on", [
        [N("1.  Why can a hold violation not be fixed by slowing the clock?")],
        [N("2.  A flop has a negative hold time of −0.02 ns. Is that good or bad for you?")],
        [N("3.  Skew of +0.3 ns on the capture clock: what happens to setup slack, and to "
           "hold slack?")],
        [N("4.  Your WNS is −0.05 ns and your TNS is −0.05 ns. What kind of problem is "
           "this?")],
        [N("5.  Why is setup signed off at the slow corner and hold at the fast corner?")],
    ], color=GREEN, fill="EEF7F1", bar="2A9D5C")

    w.page_break()
