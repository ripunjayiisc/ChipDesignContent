# -*- coding: utf-8 -*-
"""Workbook front matter + Part 1: Boolean algebra and logic gates."""
from wbkit import *
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL


def B(t, d=None, **kw):
    kw.update(d or {}); kw["b"] = True; return (t, kw)


def N(t, d=None, **kw):
    kw.update(d or {}); return (t, kw)


def I(t, d=None, **kw):
    kw.update(d or {}); kw["i"] = True; return (t, kw)


def M(t, d=None, **kw):
    kw.update(d or {}); kw["f"] = MONOF; return (t, kw)


def build(w):
    # ------------------------------------------------------------ cover
    w.para([N("CHIP DESIGN ASSOCIATE  ·  O-LEVEL ‘CHIP DESIGN’",
              {"b": True, "s": 11, "c": TEAL})], space_after=2)
    p = w.d.add_paragraph()
    r = p.add_run("Module 2 — Topic 3")
    r.font.name = HEADF; r.font.size = Pt(15); r.font.bold = True; r.font.color.rgb = SLATE
    p.paragraph_format.space_after = Pt(2)
    p = w.d.add_paragraph()
    r = p.add_run("Digital Logic Design Principles")
    r.font.name = HEADF; r.font.size = Pt(25); r.font.bold = True; r.font.color.rgb = NAVY
    p.paragraph_format.space_after = Pt(4)
    p = w.d.add_paragraph()
    r = p.add_run("Tutorial & Practice Workbook")
    r.font.name = HEADF; r.font.size = Pt(16); r.font.color.rgb = AMBER
    p.paragraph_format.space_after = Pt(10)
    w.para([N("A self-study companion to the Topic 3 slide deck. It elaborates every concept in the "
              "deck, adds the derivations the slides only quote, walks through four guided "
              "tutorials at the keyboard, and ends with 46 graded exercises and full worked "
              "solutions. Nothing here requires you to look anything up elsewhere.", {"s": 10.5})])
    w.para([N("NOS: NIE/ELE/N0102  ·  Module 2, Topic 3 (4 hours + lab)  ·  "
              "Syllabus: combinational and sequential logic design concepts; Boolean algebra and "
              "logic gates; flip-flops, registers and state machines.",
              {"s": 9, "c": SLATE, "i": True})])

    w.callout("What's inside", [
        [B("Part 1  "), N("Boolean algebra and logic gates — binary, gates, laws, De Morgan, "
                          "canonical forms, Karnaugh maps")],
        [B("Part 2  "), N("Combinational logic — design procedure, adders, MUX, decoders, "
                          "comparators, ALU, hazards, timing")],
        [B("Part 3  "), N("Sequential logic — clocks, latches, flip-flops, timing, metastability, "
                          "registers, counters, state machines")],
        [B("Part 4  "), N("Guided tutorials T1–T4 — Logisim-Evolution, Icarus Verilog, GTKWave, Yosys")],
        [B("Part 5  "), N("46 practice exercises, graded from recall to design")],
        [B("Part 6  "), N("Full worked solutions to every exercise")],
        [B("Part 7  "), N("Glossary, formula sheet and further reading")],
    ], color=TEAL)

    w.callout("How to use this workbook", [
        N("Read a section of Part 1–3 immediately after the matching slides, while the diagram is "
          "still fresh. Do the tutorial in Part 4 at a keyboard, not on paper — the whole point is "
          "that the tool agrees with your prediction. Attempt the Part 5 exercises before opening "
          "Part 6; an exercise you have read the answer to teaches you nothing."),
        [B("Symbols used throughout: "), M("A' "), N("or "), M("Ā "), N("means NOT A;  "),
         M("A·B "), N("or "), M("AB "), N("means A AND B;  "), M("A+B "),
         N("means A OR B;  "), M("A⊕B "), N("means A XOR B. Q⁺ means 'the next value of Q'.")],
    ], color=AMBER, fill="FFF7EC", bar="C77514")

    w.page_break()

    # ============================================================ PART 1
    w.h1("Part 1 · Boolean Algebra and Logic Gates")
    w.para([N("This part builds the mathematics that every digital circuit obeys, and connects it "
              "to the physical gates that implement it. Work through it in order — each section "
              "uses the one before.")])

    # ---------------------------------------------------------- 1.1
    w.h2("1.1  Why digital? The two-valued abstraction")
    w.para([N("A digital circuit is an analog circuit that we have agreed to misread. The wire "
              "genuinely carries a continuous voltage; we simply declare two bands and design so "
              "that a signal never rests between them.")])
    w.image("digital_abstraction", 6.5,
            "Figure 1.1 — a continuous voltage, read as one of two symbols")
    w.h3("The four threshold voltages")
    w.table(["Symbol", "Name", "Meaning"],
            [["V_OL", "output low", "the highest voltage a driver produces when outputting 0"],
             ["V_OH", "output high", "the lowest voltage a driver produces when outputting 1"],
             ["V_IL", "input low", "the highest voltage a receiver still reads as 0"],
             ["V_IH", "input high", "the lowest voltage a receiver still reads as 1"]],
            [0.9, 1.5, 4.3], bold_cols=(0,), align_center=False)
    w.para([B("Noise margins. "),
            N("NM_L = V_IL − V_OL and NM_H = V_OH − V_IH. These are how much interference a signal "
              "can pick up and still be read correctly. For a 3.3 V CMOS family with V_OL = 0.4, "
              "V_IL = 0.8, V_IH = 2.0, V_OH = 2.4, the margins are 0.4 V both ways.")])
    w.callout("Why this matters more than it looks",
              ["Digital logic is RESTORATIVE: every gate outputs a clean, full-swing level "
               "regardless of how degraded its input was, provided the input stayed outside the "
               "forbidden band. Noise, ageing, temperature drift and IR drop are therefore thrown "
               "away at every stage instead of accumulating. That single property is what makes a "
               "signal able to cross a chip through thousands of gates and arrive intact — and it "
               "is why analog design is hard and digital design scales."],
              color=TEAL)

    # ---------------------------------------------------------- 1.2
    w.h2("1.2  Number systems and codes")
    w.para([N("Hardware stores bits. What those bits MEAN is a convention you choose and must "
              "apply consistently. The same eight bits 1011 0100 are 180 unsigned, −76 signed, "
              "0xB4 in hex, or part of a BCD pair.")])
    w.image("number_systems", 6.5, "Figure 1.2 — one number, four representations")

    w.h3("Positional notation")
    w.para([N("In any base b, the value of a digit string is Σ dᵢ · bⁱ. In binary the weights are "
              "powers of two, so 1011 0100 = 128 + 32 + 16 + 4 = 180.")])
    w.h3("Binary ↔ hexadecimal")
    w.para([N("Because 16 = 2⁴, one hex digit is exactly one nibble. Group the bits in fours from "
              "the RIGHT and translate each group: 1011 0100 → B 4 → 0xB4. This is why register "
              "maps, opcodes and memory dumps are always written in hex — it is binary you can read.")])
    w.h3("Two's complement — the only signed format you will use")
    w.numbered([
        [N("Write the magnitude in binary: +5 in 8 bits is "), M("0000 0101"), N(".")],
        [N("Invert every bit: "), M("1111 1010"), N(".")],
        [N("Add one: "), M("1111 1011"), N(". That is −5.")],
    ])
    w.para([N("For n bits the range is −2ⁿ⁻¹ … +2ⁿ⁻¹−1 — note the asymmetry: 8 bits cover "
              "−128 … +127. The MSB is the sign bit. The great virtue is that ONE adder handles "
              "both addition and subtraction, and there is only one representation of zero.")])
    w.callout("Overflow: the trap in signed arithmetic",
              [N("Carry-out and overflow are NOT the same thing. Carry-out (C) signals unsigned "
                 "overflow. Signed overflow (V) is Cₙ ⊕ Cₙ₋₁ — the carry INTO the sign bit differs "
                 "from the carry OUT of it."),
               [B("Example: "), M("0111 1111 (+127) + 0000 0001 (+1) = 1000 0000"),
                N(" which reads as −128. C = 0 but V = 1. Adding two positives gave a negative.")],
               [B("In Verilog: "), N("declare "), M("reg signed [7:0] x;"),
                N(" or the tool assumes unsigned and your comparisons silently invert.")]],
              color=RED, fill="FDECEF", bar="C01F43")

    w.h3("Gray code")
    w.para([N("Successive values differ in exactly ONE bit: 000, 001, 011, 010, 110, 111, 101, 100. "
              "Convert binary to Gray with G = B ⊕ (B >> 1). Used for Karnaugh-map ordering, for "
              "FIFO pointers crossing clock domains, and for shaft encoders — anywhere a multi-bit "
              "value might be sampled mid-transition, because only one bit can be caught in flight.")])
    w.h3("BCD")
    w.para([N("Each decimal digit in its own 4-bit nibble: 180 → 0001 1000 0000. Six of the sixteen "
              "codes are wasted, but decimal display and exact decimal arithmetic become trivial. "
              "The mod-10 counter you build in Part 3 is a BCD counter.")])

    # ---------------------------------------------------------- 1.3
    w.h2("1.3  Boolean algebra: variables, operators, laws")
    w.para([N("Boolean algebra was formalised by George Boole in 1854 and applied to switching "
              "circuits by Claude Shannon in his 1937 master's thesis — arguably the most "
              "consequential master's thesis ever written. It is an algebra over exactly two "
              "values with three primitive operators.")])
    w.table(["Operator", "Notations", "Result is 1 when…"],
            [["AND", "A·B    A∧B    AB    A&B", "BOTH inputs are 1"],
             ["OR", "A+B    A∨B    A|B", "AT LEAST ONE input is 1"],
             ["NOT", "A'    Ā    ¬A    ~A", "the input is 0"]],
            [1.1, 2.6, 3.0], bold_cols=(0,), align_center=False)
    w.para([B("Precedence: "), N("NOT binds tightest, then AND, then OR. So A + B·C means "
              "A + (B·C). Verilog uses the same order. When in doubt, bracket it.")])

    w.h3("The truth table")
    w.para([N("A truth table lists the output for every input combination. For n inputs there are "
              "2ⁿ rows: 2 inputs → 4 rows, 4 → 16, 10 → 1 024, 16 → 65 536. That exponential is "
              "exactly why we need algebra and K-maps instead of exhaustive tables, and why "
              "exhaustive simulation stops being feasible above about 20 inputs.")])
    w.para([B("Functional equivalence. "),
            N("Two circuits are equivalent if and only if their truth tables agree on every row, "
              "however differently they are drawn. Equivalence checking — a signoff step you will "
              "meet in Topic 6 — automates precisely this comparison between your RTL and the "
              "synthesised netlist.")])

    w.h3("The laws, with their duals")
    w.table(["Law", "OR form", "AND form"],
            [["Identity", "A + 0 = A", "A · 1 = A"],
             ["Null / Dominance", "A + 1 = 1", "A · 0 = 0"],
             ["Idempotent", "A + A = A", "A · A = A"],
             ["Complement", "A + A' = 1", "A · A' = 0"],
             ["Involution", "(A')' = A", "—"],
             ["Commutative", "A + B = B + A", "A · B = B · A"],
             ["Associative", "A + (B+C) = (A+B) + C", "A · (B·C) = (A·B) · C"],
             ["Distributive", "A · (B+C) = A·B + A·C", "A + (B·C) = (A+B) · (A+C)"],
             ["Absorption", "A + A·B = A", "A · (A+B) = A"],
             ["Simplification", "A + A'·B = A + B", "A · (A'+B) = A · B"],
             ["De Morgan", "(A+B)' = A' · B'", "(A·B)' = A' + B'"],
             ["Consensus", "AB + A'C + BC = AB + A'C",
              "(A+B)(A'+C)(B+C) = (A+B)(A'+C)"]],
            [1.5, 2.6, 2.6], bold_cols=(0,), size=9, align_center=False)
    w.callout("Duality — how to halve what you memorise",
              [N("Swap every AND ↔ OR and every 0 ↔ 1 in a true Boolean statement and the result "
                 "is also true. That is why the two columns above are mirror images. Learn one "
                 "column and derive the other."),
               [B("Watch out: "), N("A + B·C = (A+B)·(A+C) has NO counterpart in ordinary "
                 "arithmetic (3 + 4×5 ≠ (3+4)×(3+5)). Verify it once on all eight rows and you "
                 "will trust it thereafter.")]],
              color=AMBER, fill="FFF7EC", bar="C77514")

    w.h3("Proving a law: perfect induction")
    w.para([N("'Perfect induction' just means 'check every row'. Because there are only 2ⁿ rows, "
              "this is a complete proof, not a sample. Here is absorption, A + A·B = A:")])
    w.table(["A", "B", "A·B", "A + A·B", "A"],
            [["0", "0", "0", "0", "0"], ["0", "1", "0", "0", "0"],
             ["1", "0", "0", "1", "1"], ["1", "1", "1", "1", "1"]],
            [0.7, 0.7, 0.9, 1.2, 0.7], bold_cols=(3, 4))
    w.para([N("Columns 4 and 5 match on every row, so the law holds. "),
            I("Intuition: if A is 1 the first term already makes the expression 1; if A is 0 "
              "neither term can be 1. B never gets a say.")])

    # ---------------------------------------------------------- 1.4
    w.h2("1.4  The logic gates")
    w.para([N("A gate is a physical circuit — in CMOS, a handful of transistors — that computes "
              "one Boolean operator. Learn to read these symbols instantly; schematics, synthesis "
              "reports and datasheets all assume you can.")])
    w.image("gate_gallery", 6.6, "Figure 1.3 — the seven gates, their expressions and truth table")

    w.h3("What a gate costs in CMOS")
    w.table(["Gate", "Transistors (static CMOS)", "Note"],
            [["Inverter (NOT)", "2", "1 PMOS + 1 NMOS"],
             ["NAND-2 / NOR-2", "4", "the natural CMOS primitives"],
             ["AND-2 / OR-2", "6", "NAND/NOR followed by an inverter"],
             ["XOR-2 / XNOR-2", "8–12", "several implementations exist"],
             ["D flip-flop", "20–28", "two latches plus clock buffering"]],
            [1.6, 2.0, 3.0], bold_cols=(0,), align_center=False)
    w.para([N("This is why standard-cell libraries are dominated by NAND, NOR and inverter cells, "
              "and why a synthesised netlist rarely contains a literal AND gate — inversion is "
              "free in CMOS, and non-inversion costs you two extra transistors.")])

    # ---------------------------------------------------------- 1.5
    w.h2("1.5  Universal gates")
    w.para([N("A gate is UNIVERSAL if any Boolean function can be built from copies of it alone. "
              "NAND and NOR are the only two-input gates with this property. AND, OR and XOR are "
              "not, because none of them can produce a NOT.")])
    w.image("universal_nand", 6.5, "Figure 1.4 — NOT, AND, OR and XOR built from NAND alone")
    w.para([B("The proof is constructive. "),
            N("Build NOT, AND and OR from the gate; then, since every Boolean function has a "
              "sum-of-products form built only from NOT, AND and OR, every function can be built. "
              "The NOR construction follows by duality.")])
    w.table(["Target", "From NAND", "From NOR"],
            [["NOT A", "A NAND A", "A NOR A"],
             ["A AND B", "(A NAND B) NAND (A NAND B)", "(A NOR A) NOR (B NOR B)"],
             ["A OR B", "(A NAND A) NAND (B NAND B)", "(A NOR B) NOR (A NOR B)"]],
            [1.3, 2.7, 2.7], bold_cols=(0,), size=9, align_center=False)

    # ---------------------------------------------------------- 1.6
    w.h2("1.6  De Morgan's theorems and bubble pushing")
    w.para([B("(A · B)' = A' + B'    and    (A + B)' = A' · B'"),
            N("  — 'break the bar, change the operator'. Both generalise to n variables.")])
    w.image("demorgan", 6.5, "Figure 1.5 — the two identities, shown as equivalent circuits")
    w.h3("Bubble pushing")
    w.para([N("Read graphically, De Morgan says a bubble (inversion) on a gate's OUTPUT can be "
              "moved to bubbles on all of its INPUTS provided you swap AND ↔ OR. This is how you "
              "convert an AND-OR network into an all-NAND network without changing behaviour, "
              "which is exactly what technology mapping does inside a synthesis tool.")])
    w.callout("Where you will actually use it: active-low signals",
              [N("A signal named rst_n, cs_n or we_n is ASSERTED when it is LOW. Specifications "
                 "written in active-low language turn into positive logic through De Morgan, and "
                 "getting it wrong is one of the most common junior-engineer bugs."),
               [B("Example: "), N("'assert busy when neither channel is idle' → "),
                M("busy = (idle_a + idle_b)' = idle_a' · idle_b'"),
                N(" — an AND of the two active-low idles, not an OR.")]],
              color=RED, fill="FDECEF", bar="C01F43")

    # ---------------------------------------------------------- 1.7
    w.h2("1.7  Canonical forms: SOP and POS")
    w.para([N("Given any truth table you can write down two exact expressions immediately, with no "
              "thinking required. They are called canonical because they are unique for a given "
              "function. They are also usually far from minimal.")])
    w.image("sop_pos", 6.5, "Figure 1.6 — the two canonical twins of one truth table")
    w.h3("Minterms and maxterms")
    w.bullets([
        [B("Minterm mᵢ "), N("is the AND term that is 1 for exactly ONE input combination. Write "
                             "each variable plain if its bit is 1, complemented if 0. For ABC = "
                             "101, m₅ = A·B'·C.")],
        [B("Maxterm Mᵢ "), N("is the OR term that is 0 for exactly ONE combination. The "
                             "complement rule is REVERSED: for ABC = 101, M₅ = A' + B + C'.")],
        [B("Relationship: "), N("Mᵢ = (mᵢ)'. A function's minterm set and maxterm set are "
                                "complementary subsets of the same 2ⁿ indices.")],
    ])
    w.h3("Writing them down")
    w.numbered([
        N("SOP: OR together one minterm for every row where Y = 1.  Y = Σm(1,3,5,7)."),
        N("POS: AND together one maxterm for every row where Y = 0.  Y = ΠM(0,2,4,6)."),
        N("Both describe the same function. In the example above, both reduce to Y = C."),
    ])
    w.para([I("Canonical is not minimal. Getting from canonical to minimal is what the next "
              "section is for.")])

    # ---------------------------------------------------------- 1.8
    w.h2("1.8  Karnaugh maps")
    w.para([N("A K-map is a truth table redrawn so that physically adjacent cells differ in "
              "exactly one variable. That turns the algebraic law X·Y + X·Y' = X into a visual "
              "one: any legal group of adjacent 1s collapses into a shorter product term.")])
    w.image("kmap_method", 6.5, "Figure 1.7 — 2-, 3- and 4-variable maps, and the grouping rules")
    w.h3("The five rules")
    w.numbered([
        N("Only ADJACENT cells may be grouped. Gray-code ordering is what guarantees that "
          "neighbours differ in exactly one variable."),
        N("Group sizes must be powers of two: 1, 2, 4, 8, 16. A group of size 2ᵏ eliminates k "
          "variables."),
        N("Make every group as LARGE as possible — a bigger group kills more variables."),
        N("Use as FEW groups as possible, but every 1 must be covered at least once. Groups may "
          "overlap."),
        N("The map WRAPS. Left edge touches right edge, top touches bottom, and the four corners "
          "are mutually adjacent. Topologically it is a torus."),
    ])
    w.h3("Vocabulary you will be examined on")
    w.bullets([
        [B("Implicant "), N("— any legal group (any product term that covers only 1s).")],
        [B("Prime implicant "), N("— an implicant that cannot be made any larger.")],
        [B("Essential prime implicant "), N("— the only prime implicant covering some particular "
                                            "minterm. It must appear in every minimal solution, "
                                            "which is why you circle these first.")],
        [B("Don't-care (X) "), N("— an input combination that cannot occur, or whose output "
                                 "nobody reads. You may treat each one as 0 or 1, whichever "
                                 "makes your groups bigger. Never cover an X for its own sake.")],
    ])
    w.h3("Worked example with don't-cares")
    w.para([N("F(A,B,C,D) = Σm(0,1,2,5,8,9,10) + d(3,7).")])
    w.image("kmap_worked", 6.5, "Figure 1.8 — grouping, including two don't-cares")
    w.numbered([
        N("Fill the map from the minterm list; mark 3 and 7 as X."),
        N("Find the essential prime implicants. m5 is covered ONLY by the group {m1,m3,m5,m7} = "
          "A'D, so A'D is essential — and note it only reached size 4 because the two X cells "
          "were absorbed."),
        N("Cover the rest with the fewest, largest groups: B'C' = {m0,m1,m8,m9} (wrapping top to "
          "bottom) and B'D' = {m0,m2,m8,m10} (the four corners)."),
        N("Read each group: keep the variables that stay constant across it, drop those that change."),
        [B("F = B'C' + B'D' + A'D"), N("  — 3 terms, 6 literals, against 7 terms and 28 literals "
                                        "for the canonical SOP.")],
        N("Verify: check that every original 1 is covered and no 0 is."),
    ])
    w.h3("Five-variable maps and beyond")
    w.para([N("A 5-variable map is drawn as two 4-variable maps side by side, one for E = 0 and "
              "one for E = 1; cells in the same position on the two maps are adjacent. Six "
              "variables needs four maps and most people stop being reliable. Past that, use the "
              "algorithms in §1.9.")])

    # ---------------------------------------------------------- 1.9
    w.h2("1.9  Beyond K-maps — what tools actually do")
    w.table(["Method", "How it works", "Scales to", "Where you meet it"],
            [["Karnaugh map", "visual grouping on a Gray-coded grid", "≤ 5–6 vars",
              "exams, whiteboards, intuition"],
             ["Quine–McCluskey", "tabular; enumerates ALL prime implicants, then solves a "
              "covering problem", "≤ ~15 vars", "the exact algorithm K-maps approximate"],
             ["Espresso", "heuristic expand / reduce / irredundant loop", "hundreds",
              "the classic two-level minimiser"],
             ["ABC / AIG rewriting", "multi-level: And-Inverter Graphs, cut rewriting, SAT "
              "sweeping", "millions", "inside Yosys and every commercial tool"],
             ["BDD", "canonical decision-diagram form", "varies wildly",
              "equivalence checking, formal verification"]],
            [1.4, 2.4, 1.0, 2.0], bold_cols=(0,), size=8.5, align_center=False)
    w.h3("Quine–McCluskey in outline")
    w.numbered([
        N("List all minterms (and don't-cares) grouped by the number of 1s in their binary code."),
        N("Compare every term in group k with every term in group k+1. Any pair differing in one "
          "bit combines; write the combined term with a dash in that position and tick both parents."),
        N("Repeat on the new list until nothing more combines. Every UNTICKED term is a prime "
          "implicant."),
        N("Build a prime-implicant chart: rows are prime implicants, columns are minterms (not "
          "don't-cares). A column with a single tick identifies an ESSENTIAL prime implicant."),
        N("Select all essentials, delete the minterms they cover, then solve the remaining "
          "covering problem — by inspection, or by Petrick's method for an exact answer."),
    ])
    w.callout("So why learn K-maps at all?",
              ["Because you must be able to read what the tool produced and judge whether it is "
               "right. A synthesis report showing 4 000 gates for a function you know needs 40 "
               "means your RTL is wrong — an unintended latch, a runaway loop, a missing "
               "don't-care, a mis-sized bus. Minimisation intuition is your only defence against "
               "silently accepting bad hardware."],
              color=GREEN, fill="EEF7F1", bar="2A9D5C")

    w.h3("What minimisation buys you")
    w.image("gate_cost", 6.5, "Figure 1.9 — the same function, before and after minimisation")
    w.para([N("Assume a 28 nm library where one gate-equivalent (GE) ≈ 0.5 µm², and a 2-input gate "
              "switching at 500 MHz burns ≈ 0.6 µW. Going from 12 GE to 5 GE saves 3.5 µm² and "
              "4.2 µW per instance. Multiply by 200 000 instances on a real SoC and that is "
              "0.7 mm² of die and 0.84 W — the difference between a product that ships and one "
              "that does not.")])
    w.para([I("Numbers are illustrative and chosen to show the method; always use your own "
              "library's datasheet.")])
    w.page_break()
