# -*- coding: utf-8 -*-
"""Styling helpers for the Chip Design Associate tutorial workbooks.

Matches Module1_Tutorial_Practice_Workbook.docx: Georgia headings in navy with
an amber rule under H1, Calibri 10.5 body, navy table headers.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY   = RGBColor(0x0E, 0x2A, 0x47)
TEAL   = RGBColor(0x1B, 0x9A, 0xAA)
AMBER  = RGBColor(0xC7, 0x75, 0x14)
GREEN  = RGBColor(0x2A, 0x9D, 0x5C)
RED    = RGBColor(0xC0, 0x1F, 0x43)
BODY   = RGBColor(0x33, 0x41, 0x4F)
SLATE  = RGBColor(0x5A, 0x6B, 0x7B)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
VIOLET = RGBColor(0x7A, 0x4F, 0xBF)

HEADF, BODYF, MONOF = "Georgia", "Calibri", "Consolas"
IMG = os.environ.get("CDA_IMG_DIR",
                     os.path.join(os.path.dirname(os.path.abspath(__file__)), "img"))


def _sub(parent, tag, **attrs):
    el = OxmlElement(tag)
    for k, v in attrs.items():
        el.set(qn(k.replace("_", ":")), v)
    parent.append(el)
    return el


def shade(cell_or_para, hexcolor):
    pr = cell_or_para._tc.get_or_add_tcPr() if hasattr(cell_or_para, "_tc") \
        else cell_or_para._p.get_or_add_pPr()
    _sub(pr, "w:shd", w_val="clear", w_color="auto", w_fill=hexcolor)


def borders(cell, color="D8DEE5", sz="4"):
    tcPr = cell._tc.get_or_add_tcPr()
    b = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement("w:" + edge)
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), sz)
        e.set(qn("w:space"), "0");    e.set(qn("w:color"), color)
        b.append(e)
    tcPr.append(b)


def no_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    b = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement("w:" + edge)
        e.set(qn("w:val"), "nil")
        b.append(e)
    tcPr.append(b)


def left_bar(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    b = OxmlElement("w:tcBorders")
    for edge, val, sz in (("top", "nil", "0"), ("bottom", "nil", "0"),
                          ("right", "nil", "0"), ("left", "single", "18")):
        e = OxmlElement("w:" + edge)
        e.set(qn("w:val"), val); e.set(qn("w:sz"), sz)
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), color)
        b.append(e)
    tcPr.append(b)


class Workbook:
    def __init__(self, footer=None):
        self.footer_text = footer or (
            "Module 2 · Topic 3 · Digital Logic Design Principles   ·   "
            "Chip Design Associate (O-Level ‘Chip Design’)   ·   NIE/ELE/N0102")
        self.d = Document()
        sec = self.d.sections[0]
        sec.page_width, sec.page_height = Inches(8.5), Inches(11)
        sec.left_margin = sec.right_margin = Inches(0.85)
        sec.top_margin = sec.bottom_margin = Inches(0.85)
        self._styles()
        self._footer()

    # ------------------------------------------------------------- styles
    def _styles(self):
        st = self.d.styles
        n = st["Normal"]
        n.font.name = BODYF
        n.font.size = Pt(10.5)
        n.font.color.rgb = BODY
        n.paragraph_format.space_after = Pt(6)
        n.paragraph_format.line_spacing = 1.12
        specs = [("Heading 1", 17, NAVY, True, 14, 7),
                 ("Heading 2", 13.5, NAVY, True, 11, 5),
                 ("Heading 3", 11.5, TEAL, True, 9, 3),
                 ("Heading 4", 10.5, AMBER, True, 7, 2)]
        for name, size, col, bold, before, after in specs:
            s = st[name]
            s.font.name = HEADF
            s.font.size = Pt(size)
            s.font.bold = bold
            s.font.color.rgb = col
            s.paragraph_format.space_before = Pt(before)
            s.paragraph_format.space_after = Pt(after)
            s.paragraph_format.keep_with_next = True
        # amber rule under H1
        h1 = st["Heading 1"]._element.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6")
        bot.set(qn("w:space"), "6");    bot.set(qn("w:color"), "C77514")
        pbdr.append(bot); h1.append(pbdr)

    def _footer(self):
        f = self.d.sections[0].footer
        p = f.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(self.footer_text)
        r.font.size = Pt(7.5)
        r.font.color.rgb = SLATE
        r.font.name = BODYF

    # ------------------------------------------------------------- blocks
    def para(self, runs, size=10.5, align=None, space_after=6, style=None,
             indent=0):
        p = self.d.add_paragraph(style=style)
        if align:
            p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        if indent:
            p.paragraph_format.left_indent = Inches(indent)
        if isinstance(runs, str):
            runs = [(runs, {})]
        for item in runs:
            txt, kw = item if isinstance(item, tuple) else (item, {})
            r = p.add_run(txt)
            r.font.name = kw.get("f", BODYF)
            r.font.size = Pt(kw.get("s", size))
            r.font.bold = kw.get("b", False)
            r.font.italic = kw.get("i", False)
            r.font.color.rgb = kw.get("c", BODY)
        return p

    def h1(self, t): return self.d.add_paragraph(t, style="Heading 1")
    def h2(self, t): return self.d.add_paragraph(t, style="Heading 2")
    def h3(self, t): return self.d.add_paragraph(t, style="Heading 3")
    def h4(self, t): return self.d.add_paragraph(t, style="Heading 4")

    def bullets(self, items, size=10.5, indent=0.22):
        for it in items:
            p = self.d.add_paragraph()
            p.paragraph_format.left_indent = Inches(indent)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run("▪  ")
            r.font.color.rgb = TEAL; r.font.size = Pt(size * 0.85); r.font.bold = True
            if isinstance(it, str):
                it = [(it, {})]
            for item in it:
                txt, kw = item if isinstance(item, tuple) else (item, {})
                rr = p.add_run(txt)
                rr.font.name = kw.get("f", BODYF)
                rr.font.size = Pt(kw.get("s", size))
                rr.font.bold = kw.get("b", False)
                rr.font.italic = kw.get("i", False)
                rr.font.color.rgb = kw.get("c", BODY)

    def numbered(self, items, size=10.5, indent=0.22, start=1):
        for i, it in enumerate(items, start):
            p = self.d.add_paragraph()
            p.paragraph_format.left_indent = Inches(indent)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run("%d.  " % i)
            r.font.color.rgb = TEAL; r.font.size = Pt(size); r.font.bold = True
            if isinstance(it, str):
                it = [(it, {})]
            for item in it:
                txt, kw = item if isinstance(item, tuple) else (item, {})
                rr = p.add_run(txt)
                rr.font.name = kw.get("f", BODYF)
                rr.font.size = Pt(kw.get("s", size))
                rr.font.bold = kw.get("b", False)
                rr.font.italic = kw.get("i", False)
                rr.font.color.rgb = kw.get("c", BODY)

    def callout(self, title, body, color=TEAL, fill="F4F8FB", bar="1B9AAA"):
        t = self.d.add_table(rows=1, cols=1)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        c = t.cell(0, 0)
        shade(c, fill)
        left_bar(c, bar)
        c.paragraphs[0].text = ""
        if title:
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(title)
            r.font.name = HEADF; r.font.size = Pt(10.5); r.font.bold = True
            r.font.color.rgb = color
            p2 = c.add_paragraph()
        else:
            p2 = c.paragraphs[0]
        p2.paragraph_format.space_after = Pt(2)
        if isinstance(body, str):
            body = [body]
        for i, line in enumerate(body):
            p = p2 if i == 0 else c.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            if isinstance(line, str):
                line = [(line, {})]
            for item in line:
                txt, kw = item if isinstance(item, tuple) else (item, {})
                r = p.add_run(txt)
                r.font.name = kw.get("f", BODYF)
                r.font.size = Pt(kw.get("s", 10))
                r.font.bold = kw.get("b", False)
                r.font.italic = kw.get("i", False)
                r.font.color.rgb = kw.get("c", BODY)
        self.para("", space_after=2)
        return t

    def code(self, lines, caption=None, size=8.8):
        if caption:
            p = self.d.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(caption)
            r.font.name = BODYF; r.font.size = Pt(9); r.font.bold = True
            r.font.color.rgb = TEAL
        t = self.d.add_table(rows=1, cols=1)
        c = t.cell(0, 0)
        shade(c, "11212F")
        no_borders(c)
        for i, ln in enumerate(lines):
            p = c.paragraphs[0] if i == 0 else c.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            r = p.add_run(ln if ln else " ")
            r.font.name = MONOF
            r.font.size = Pt(size)
            r.font.color.rgb = RGBColor(0xDC, 0xE6, 0xF0)
        self.para("", space_after=2)
        return t

    def table(self, headers, rows, widths=None, size=9.5, head_fill="0E2A47",
              bold_cols=(), align_center=True):
        t = self.d.add_table(rows=1 + len(rows), cols=len(headers))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for j, h in enumerate(headers):
            c = t.cell(0, j)
            shade(c, head_fill); borders(c, "0E2A47", "4")
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(h))
            r.font.bold = True; r.font.size = Pt(size); r.font.color.rgb = WHITE
            r.font.name = BODYF
        for i, row in enumerate(rows, 1):
            for j, cell in enumerate(row):
                c = t.cell(i, j)
                borders(c)
                if i % 2 == 0:
                    shade(c, "F4F8FB")
                p = c.paragraphs[0]
                p.paragraph_format.space_after = Pt(1)
                if align_center and j > 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(str(cell))
                r.font.size = Pt(size)
                r.font.name = MONOF if str(cell).strip().startswith("`") else BODYF
                if r.font.name == MONOF:
                    r.text = str(cell).strip().strip("`")
                r.font.bold = j in bold_cols
                r.font.color.rgb = NAVY if j in bold_cols else BODY
        if widths:
            for j, w in enumerate(widths):
                for row in t.rows:
                    row.cells[j].width = Inches(w)
        self.para("", space_after=4)
        return t

    def image(self, name, width=6.6, caption=None):
        self.d.add_picture(os.path.join(os.environ.get("CDA_IMG_DIR", IMG),
                                        name + ".png"), width=Inches(width))
        self.d.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            p = self.d.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(8)
            r = p.add_run(caption)
            r.font.size = Pt(8.5); r.font.italic = True; r.font.color.rgb = SLATE

    def page_break(self):
        self.d.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    def save(self, path):
        self.d.save(path)
        print("saved", path)
